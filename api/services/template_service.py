"""
Template service for managing NDA templates and rendering them with data
"""
import uuid
import logging
import tempfile
import os
from typing import Dict, Any, Optional, List
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt
import re

from api.db import get_db_session
from api.db.schema import NDATemplate
from api.services.service_registry import get_storage_service

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("docx2pdf not available. PDF conversion will use fallback method.")

logger = logging.getLogger(__name__)

TEMPLATE_BUCKET = "nda-templates"


class TemplateService:
    """Service for managing and rendering NDA templates"""

    def __init__(self):
        self.storage = get_storage_service()
        self._ensure_template_bucket()

    def _ensure_template_bucket(self):
        """Ensure template bucket exists"""
        try:
            if not self.storage.file_exists(TEMPLATE_BUCKET, ".bucket_check"):
                # Try to create bucket by uploading a dummy file
                try:
                    self.storage.upload_file(
                        TEMPLATE_BUCKET,
                        ".bucket_check",
                        b"",
                        "text/plain"
                    )
                except Exception:
                    # Bucket might already exist, that's fine
                    pass
        except Exception as e:
            logger.warning(f"Could not ensure template bucket exists: {e}")

    def create_template(
        self,
        name: str,
        description: Optional[str],
        file_data: bytes,
        created_by: Optional[str] = None,
        template_key: Optional[str] = None,
        change_notes: Optional[str] = None,
    ) -> NDATemplate:
        """
        Create a new template or new version of existing template
        
        Args:
            name: Template name
            description: Template description
            file_data: Template file bytes (DOCX)
            created_by: User ID who created the template
            template_key: Unique key for template (auto-generated from name if not provided)
            change_notes: Notes about changes in this version
            
        Returns:
            Created NDATemplate instance
        """
        # Validate it's a DOCX file
        if not file_data.startswith(b'PK'):
            raise ValueError("Template must be a valid DOCX file")
        
        # Generate template_key from name if not provided
        if not template_key:
            import re
            template_key = re.sub(r'[^a-zA-Z0-9]+', '-', name.lower()).strip('-')
            if not template_key:
                template_key = f"template-{uuid.uuid4().hex[:8]}"
        
        db = get_db_session()
        try:
            # Check if template_key exists - if so, create new version
            existing_templates = db.query(NDATemplate).filter(
                NDATemplate.template_key == template_key
            ).order_by(NDATemplate.version.desc()).all()
            
            if existing_templates:
                # Create new version
                latest_version = existing_templates[0]
                new_version = latest_version.version + 1
                
                # Mark all previous versions as not current
                db.query(NDATemplate).filter(
                    NDATemplate.template_key == template_key
                ).update({"is_current": False})
                
                logger.info(f"Creating version {new_version} of template {template_key}")
            else:
                # First version
                new_version = 1
                logger.info(f"Creating new template {template_key} version 1")
            
            # Generate unique file path with version
            template_id = str(uuid.uuid4())
            file_path = f"{template_key}/v{new_version}/{name}.docx"
            
            # Upload template to storage
            storage_path = self.storage.upload_file(
                TEMPLATE_BUCKET,
                file_path,
                file_data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            created_by_uuid = None
            if created_by:
                try:
                    created_by_uuid = uuid.UUID(created_by)
                except ValueError:
                    logger.warning(f"Invalid user ID: {created_by}")
            
            template = NDATemplate(
                name=name,
                description=description,
                file_path=storage_path,
                template_key=template_key,
                version=new_version,
                is_active=True,
                is_current=True,
                change_notes=change_notes,
                created_by=created_by_uuid,
            )
            
            db.add(template)
            db.commit()
            db.refresh(template)
            
            logger.info(f"Created template: {name} ({template.id})")
            return template
        except Exception as e:
            logger.error(f"Failed to create template: {e}")
            db.rollback()
            # Clean up uploaded file
            try:
                self.storage.delete_file(TEMPLATE_BUCKET, file_path)
            except Exception:
                pass
            raise
        finally:
            db.close()

    def get_template(self, template_id: str, version: Optional[int] = None) -> Optional[NDATemplate]:
        """
        Get template by ID, optionally by specific version
        
        Args:
            template_id: Template ID or template_key
            version: Specific version number (None = get current version)
            
        Returns:
            NDATemplate instance or None
        """
        db = get_db_session()
        try:
            # Try as UUID first
            try:
                template_uuid = uuid.UUID(template_id)
                query = db.query(NDATemplate).filter(NDATemplate.id == template_uuid)
            except ValueError:
                # Assume it's a template_key
                query = db.query(NDATemplate).filter(NDATemplate.template_key == template_id)
            
            if version:
                query = query.filter(NDATemplate.version == version)
            else:
                # Get current version
                query = query.filter(NDATemplate.is_current == True)
            
            template = query.first()
            return template
        except Exception as e:
            logger.error(f"Error getting template: {e}")
            return None
        finally:
            db.close()

    def list_templates(self, active_only: bool = True, current_only: bool = True) -> List[NDATemplate]:
        """
        List all templates
        
        Args:
            active_only: Only return active templates
            current_only: Only return current versions (latest version of each template_key)
            
        Returns:
            List of NDATemplate instances
        """
        db = get_db_session()
        try:
            query = db.query(NDATemplate)
            if active_only:
                query = query.filter(NDATemplate.is_active == True)
            if current_only:
                query = query.filter(NDATemplate.is_current == True)
            return query.order_by(NDATemplate.template_key, NDATemplate.version.desc()).all()
        finally:
            db.close()

    def list_template_versions(self, template_key: str) -> List[NDATemplate]:
        """
        List all versions of a template
        
        Args:
            template_key: Template key to get versions for
            
        Returns:
            List of NDATemplate instances ordered by version (newest first)
        """
        db = get_db_session()
        try:
            return db.query(NDATemplate).filter(
                NDATemplate.template_key == template_key
            ).order_by(NDATemplate.version.desc()).all()
        finally:
            db.close()

    def set_current_version(self, template_key: str, version: int) -> Optional[NDATemplate]:
        """
        Set a specific version as the current/default version
        
        Args:
            template_key: Template key
            version: Version number to set as current
            
        Returns:
            Updated NDATemplate instance or None
        """
        db = get_db_session()
        try:
            # Get the template version
            template = db.query(NDATemplate).filter(
                NDATemplate.template_key == template_key,
                NDATemplate.version == version
            ).first()
            
            if not template:
                return None
            
            # Mark all versions as not current
            db.query(NDATemplate).filter(
                NDATemplate.template_key == template_key
            ).update({"is_current": False})
            
            # Mark this version as current
            template.is_current = True
            db.commit()
            db.refresh(template)
            
            logger.info(f"Set template {template_key} version {version} as current")
            return template
        except Exception as e:
            logger.error(f"Failed to set current version: {e}")
            db.rollback()
            return None
        finally:
            db.close()

    def update_template(
        self,
        template_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        is_active: Optional[bool] = None,
    ) -> Optional[NDATemplate]:
        """Update template metadata"""
        db = get_db_session()
        try:
            template_uuid = uuid.UUID(template_id)
            template = db.query(NDATemplate).filter(NDATemplate.id == template_uuid).first()
            
            if not template:
                return None
            
            if name is not None:
                template.name = name
            if description is not None:
                template.description = description
            if is_active is not None:
                template.is_active = is_active
            
            db.commit()
            db.refresh(template)
            return template
        except ValueError:
            return None
        except Exception as e:
            logger.error(f"Failed to update template: {e}")
            db.rollback()
            return None
        finally:
            db.close()

    def delete_template(self, template_id: str, version: Optional[int] = None) -> bool:
        """
        Delete template or specific version
        
        Args:
            template_id: Template ID or template_key
            version: Specific version to delete (None = delete all versions)
            
        Returns:
            True if deleted successfully, False otherwise
        """
        db = get_db_session()
        try:
            # Get template(s) to delete
            try:
                template_uuid = uuid.UUID(template_id)
                query = db.query(NDATemplate).filter(NDATemplate.id == template_uuid)
            except ValueError:
                # Assume it's a template_key
                query = db.query(NDATemplate).filter(NDATemplate.template_key == template_id)
            
            if version:
                query = query.filter(NDATemplate.version == version)
            
            templates = query.all()
            
            if not templates:
                return False
            
            # Delete files from storage and database records
            deleted_count = 0
            for template in templates:
                # Delete file from storage
                try:
                    # Parse file_path (format: bucket/object_name)
                    file_path = template.file_path
                    if '/' in file_path:
                        parts = file_path.split('/', 1)
                        bucket = parts[0]
                        object_name = parts[1] if len(parts) > 1 else file_path
                        
                        # Normalize bucket name
                        if bucket == "test":
                            bucket = TEMPLATE_BUCKET
                        
                        if bucket == TEMPLATE_BUCKET or bucket == "nda-templates":
                            self.storage.delete_file(bucket, object_name)
                except Exception as e:
                    logger.warning(f"Failed to delete template file for {template.id}: {e}")
                
                # Delete from database
                db.delete(template)
                deleted_count += 1
            
            db.commit()
            
            logger.info(f"Deleted {deleted_count} template(s): {template_id}" + (f" version {version}" if version else ""))
            return True
        except ValueError:
            return False
        except Exception as e:
            logger.error(f"Failed to delete template: {e}", exc_info=True)
            db.rollback()
            return False
        finally:
            db.close()
    
    def delete_template_version(self, template_key: str, version: int) -> bool:
        """
        Delete a specific version of a template
        
        Args:
            template_key: Template key
            version: Version number to delete
            
        Returns:
            True if deleted successfully, False otherwise
        """
        return self.delete_template(template_key, version=version)

    def render_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        version: Optional[int] = None,
    ) -> bytes:
        """
        Render template with provided data
        
        Args:
            template_id: Template ID or template_key
            data: Dictionary with template variables to fill
            version: Specific version to use (None = use current version)
            
        Returns:
            Rendered DOCX file as bytes
        """
        template = self.get_template(template_id, version=version)
        if not template:
            raise ValueError(f"Template not found: {template_id}" + (f" version {version}" if version else ""))
        
        if not template.is_active:
            raise ValueError(f"Template is not active: {template_id}")
        
        # Get raw template file
        template_bytes = self.get_template_file(template_id, version=version)
        
        # Load DOCX document
        doc = Document(BytesIO(template_bytes))
        
        # Prepare replacement data
        replacements = self._prepare_replacements(data)
        
        # Replace placeholders in document
        self._replace_placeholders(doc, replacements)
        
        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
    
    def get_template_file(self, template_id: str, version: Optional[int] = None) -> bytes:
        """
        Get raw template file bytes (without rendering)
        
        Args:
            template_id: Template ID or template_key
            version: Specific version number (None = use current version)
            
        Returns:
            Raw template file bytes
        """
        template = self.get_template(template_id, version=version)
        if not template:
            raise ValueError(f"Template not found: {template_id}" + (f" version {version}" if version else ""))
        
        # Download template file
        file_path = template.file_path
        
        # Parse file_path - upload_file returns "bucket/object_name" format
        # But we also need to handle legacy formats
        if '/' in file_path:
            parts = file_path.split('/', 1)
            bucket = parts[0]
            object_name = parts[1] if len(parts) > 1 else file_path
            
            # Handle legacy "test/" prefix or normalize bucket name
            if bucket == "test":
                bucket = TEMPLATE_BUCKET
            elif bucket not in [TEMPLATE_BUCKET, "nda-templates"]:
                # If bucket doesn't match expected, assume it's part of object_name
                # This handles legacy formats where bucket wasn't included
                bucket = TEMPLATE_BUCKET
                object_name = file_path
        else:
            # No slash - entire path is object_name
            bucket = TEMPLATE_BUCKET
            object_name = file_path
        
        try:
            template_bytes = self.storage.download_file(bucket, object_name)
        except Exception as e:
            logger.error(f"Failed to download template file from {bucket}/{object_name} (file_path={file_path}): {e}")
            raise ValueError(f"Failed to retrieve template file: {str(e)}")
        
        return template_bytes
    
    def convert_template_to_pdf(self, template_id: str, version: Optional[int] = None) -> bytes:
        """
        Convert template DOCX to PDF
        
        Args:
            template_id: Template ID or template_key
            version: Specific version number (None = use current version)
            
        Returns:
            PDF file bytes
        """
        template = self.get_template(template_id, version=version)
        if not template:
            raise ValueError(f"Template not found: {template_id}" + (f" version {version}" if version else ""))
        
        # Get DOCX file
        docx_bytes = self.get_template_file(template_id, version=version)
        
        # Convert DOCX to PDF
        if DOCX2PDF_AVAILABLE:
            try:
                # Create temporary files
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                    docx_file.write(docx_bytes)
                    docx_path = docx_file.name
                
                pdf_path = docx_path.replace('.docx', '.pdf')
                
                try:
                    # Convert using docx2pdf (requires LibreOffice or MS Word)
                    convert(docx_path, pdf_path)
                    
                    # Read PDF bytes
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_bytes = pdf_file.read()
                    
                    return pdf_bytes
                finally:
                    # Clean up temp files
                    try:
                        os.unlink(docx_path)
                    except:
                        pass
                    try:
                        os.unlink(pdf_path)
                    except:
                        pass
            except Exception as e:
                logger.error(f"Failed to convert DOCX to PDF using docx2pdf: {e}")
                # Fall through to fallback method
        
        # Fallback: Use reportlab to create a simple PDF (basic conversion)
        # This is a simplified fallback - for better results, use LibreOffice
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            # Parse DOCX and create PDF
            doc = Document(BytesIO(docx_bytes))
            
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            y = height - 50
            margin = 50
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Simple text rendering (doesn't preserve formatting)
                    text = para.text[:100]  # Limit line length
                    if y < margin:
                        p.showPage()
                        y = height - 50
                    p.drawString(margin, y, text)
                    y -= 20
            
            p.save()
            buffer.seek(0)
            return buffer.read()
        except Exception as e:
            logger.error(f"Failed to convert DOCX to PDF using fallback: {e}")
            raise ValueError(f"Failed to convert template to PDF: {str(e)}")
    
    def render_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        version: Optional[int] = None,
    ) -> bytes:
        """
        Render template with provided data
        
        Args:
            template_id: Template ID or template_key
            data: Dictionary with template variables to fill
            version: Specific version to use (None = use current version)
            
        Returns:
            Rendered DOCX file as bytes
        """
        template = self.get_template(template_id, version=version)
        if not template:
            raise ValueError(f"Template not found: {template_id}" + (f" version {version}" if version else ""))
        
        if not template.is_active:
            raise ValueError(f"Template is not active: {template_id}")
        
        # Get raw template file
        template_bytes = self.get_template_file(template_id, version=version)
        
        # Load DOCX document
        doc = Document(BytesIO(template_bytes))
        
        # Prepare replacement data
        replacements = self._prepare_replacements(data)
        
        # Replace placeholders in document
        self._replace_placeholders(doc, replacements)
        
        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _prepare_replacements(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Prepare replacement dictionary from data"""
        replacements = {}
        
        # Standard fields
        replacements['{counterparty_name}'] = str(data.get('counterparty_name', ''))
        replacements['{effective_date}'] = self._format_date(data.get('effective_date'))
        replacements['{term_months}'] = str(data.get('term_months', ''))
        replacements['{survival_months}'] = str(data.get('survival_months', ''))
        replacements['{governing_law}'] = str(data.get('governing_law', ''))
        replacements['{disclosing_party}'] = str(data.get('disclosing_party', ''))
        replacements['{receiving_party}'] = str(data.get('receiving_party', ''))
        
        # Date formatting
        if data.get('effective_date'):
            effective_date = data['effective_date']
            if isinstance(effective_date, date):
                replacements['{effective_date_formatted}'] = effective_date.strftime('%B %d, %Y')
                replacements['{effective_date_short}'] = effective_date.strftime('%m/%d/%Y')
            else:
                replacements['{effective_date_formatted}'] = str(effective_date)
                replacements['{effective_date_short}'] = str(effective_date)
        
        # Add any custom fields from data
        for key, value in data.items():
            placeholder = f"{{{key}}}"
            if placeholder not in replacements:
                replacements[placeholder] = str(value) if value is not None else ''
        
        return replacements

    def _format_date(self, date_value: Any) -> str:
        """Format date value for template"""
        if date_value is None:
            return ''
        if isinstance(date_value, date):
            return date_value.strftime('%B %d, %Y')
        return str(date_value)

    def _replace_placeholders(self, doc: Document, replacements: Dict[str, str]):
        """Replace placeholders in document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in headers and footers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                for paragraph in header.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for placeholder, value in replacements.items():
                                    if placeholder in paragraph.text:
                                        paragraph.text = paragraph.text.replace(placeholder, value)

    def get_template_variables(self, template_id: str) -> List[str]:
        """
        Extract template variables from a template
        Returns list of variable names found in template
        """
        template = self.get_template(template_id)
        if not template:
            return []
        
        # Download template file
        file_path = template.file_path
        if '/' in file_path:
            parts = file_path.split('/', 1)
            bucket = parts[0] if parts[0] else TEMPLATE_BUCKET
            object_name = parts[1] if len(parts) > 1 else file_path
        else:
            bucket = TEMPLATE_BUCKET
            object_name = file_path
        
        try:
            template_bytes = self.storage.download_file(bucket, object_name)
            doc = Document(BytesIO(template_bytes))
            
            # Extract all placeholders using regex
            variables = set()
            pattern = r'\{([^}]+)\}'
            
            # Check paragraphs
            for paragraph in doc.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                variables.update(matches)
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(pattern, paragraph.text)
                            variables.update(matches)
            
            return sorted(list(variables))
        except Exception as e:
            logger.error(f"Failed to extract template variables: {e}")
            return []


def create_template_service() -> TemplateService:
    """Factory function for template service"""
    return TemplateService()


