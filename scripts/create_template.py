#!/usr/bin/env python3
"""
Create a sample NDA template DOCX file
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def create_nda_template(output_path='templates/template_standard_mutual_nda.docx'):
    """Create a standard mutual NDA template with placeholder variables"""
    
    # Create a new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('MUTUAL NON-DISCLOSURE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Parties section
    doc.add_heading('1. PARTIES', level=1)
    p = doc.add_paragraph()
    p.add_run('This Mutual Non-Disclosure Agreement ("Agreement") is entered into on ')
    p.add_run('{{effective_date}}').bold = True
    p.add_run(' between:')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Disclosing Party: ').bold = True
    p.add_run('{{disclosing_party}}')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Receiving Party: ').bold = True
    p.add_run('{{receiving_party}}')
    doc.add_paragraph()
    
    # Purpose section
    doc.add_heading('2. PURPOSE', level=1)
    doc.add_paragraph('The parties wish to engage in discussions regarding potential business opportunities and may disclose confidential information to each other in connection with such discussions.')
    
    # Definition of Confidential Information
    doc.add_heading('3. DEFINITION OF CONFIDENTIAL INFORMATION', level=1)
    doc.add_paragraph('For purposes of this Agreement, "Confidential Information" means all non-public, proprietary, or confidential information disclosed by one party (the "Disclosing Party") to the other party (the "Receiving Party"), whether orally, in writing, or in any other form, including but not limited to:')
    doc.add_paragraph('(a) Technical data, know-how, research, product plans, products, services, customers, customer lists, markets, software, developments, inventions, processes, formulas, technology, designs, drawings, engineering, hardware configuration information, marketing, finances, or other business information;', style='List Bullet')
    doc.add_paragraph('(b) Any information that the Disclosing Party designates as confidential or proprietary, or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure.', style='List Bullet')
    
    # Exclusions
    doc.add_heading('4. EXCLUSIONS', level=1)
    doc.add_paragraph('Confidential Information does not include information that:')
    doc.add_paragraph('(a) Is or becomes publicly known through no breach of this Agreement by the Receiving Party;', style='List Bullet')
    doc.add_paragraph('(b) Was rightfully known by the Receiving Party prior to disclosure;', style='List Bullet')
    doc.add_paragraph('(c) Is rightfully received by the Receiving Party from a third party without breach of any confidentiality obligation;', style='List Bullet')
    doc.add_paragraph('(d) Is independently developed by the Receiving Party without use of or reference to the Confidential Information.', style='List Bullet')
    
    # Obligations
    doc.add_heading('5. OBLIGATIONS', level=1)
    doc.add_paragraph('Each Receiving Party agrees to:')
    doc.add_paragraph('(a) Hold and maintain the Confidential Information in strict confidence;', style='List Bullet')
    doc.add_paragraph('(b) Not disclose the Confidential Information to any third parties without the prior written consent of the Disclosing Party;', style='List Bullet')
    doc.add_paragraph('(c) Use the Confidential Information solely for the purpose of evaluating potential business opportunities between the parties;', style='List Bullet')
    doc.add_paragraph('(d) Take reasonable precautions to protect the confidentiality of the Confidential Information, using at least the same degree of care it uses to protect its own confidential information.', style='List Bullet')
    
    # Term
    doc.add_heading('6. TERM', level=1)
    p = doc.add_paragraph()
    p.add_run('This Agreement shall remain in effect for a period of ')
    p.add_run('{{term_months}}').bold = True
    p.add_run(' months from the effective date, unless terminated earlier by mutual written consent of the parties.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Survival Period: ').bold = True
    p.add_run('The obligations of confidentiality shall survive termination of this Agreement and continue for ')
    p.add_run('{{survival_months}}').bold = True
    p.add_run(' months after termination.')
    
    # Return of Materials
    doc.add_heading('7. RETURN OF MATERIALS', level=1)
    doc.add_paragraph('Upon termination of this Agreement or upon written request by the Disclosing Party, the Receiving Party shall promptly return or destroy all documents, materials, and other tangible manifestations of Confidential Information and all copies thereof.')
    
    # No License
    doc.add_heading('8. NO LICENSE', level=1)
    doc.add_paragraph('Nothing in this Agreement shall be construed as granting any rights, by license or otherwise, to any Confidential Information, except as expressly set forth herein.')
    
    # Remedies
    doc.add_heading('9. REMEDIES', level=1)
    doc.add_paragraph('Each party acknowledges that any breach of this Agreement may cause irreparable harm to the Disclosing Party, and that monetary damages may be inadequate to compensate for such breach. Therefore, the Disclosing Party shall be entitled to seek injunctive relief and other equitable remedies in addition to any other remedies available at law or in equity.')
    
    # Governing Law
    doc.add_heading('10. GOVERNING LAW', level=1)
    p = doc.add_paragraph()
    p.add_run('This Agreement shall be governed by and construed in accordance with the laws of ')
    p.add_run('{{governing_law}}').bold = True
    p.add_run('.')
    
    # Entire Agreement
    doc.add_heading('11. ENTIRE AGREEMENT', level=1)
    doc.add_paragraph('This Agreement constitutes the entire agreement between the parties concerning the subject matter hereof and supersedes all prior agreements, understandings, negotiations, and discussions, whether oral or written.')
    
    # Signatures section
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Disclosing Party signature
    p = doc.add_paragraph()
    p.add_run('DISCLOSING PARTY:').bold = True
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    p = doc.add_paragraph()
    p.add_run('{{disclosing_party}}')
    doc.add_paragraph()
    doc.add_paragraph('By: _________________________')
    doc.add_paragraph('Name:')
    doc.add_paragraph('Title:')
    doc.add_paragraph('Date:')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Receiving Party signature
    p = doc.add_paragraph()
    p.add_run('RECEIVING PARTY:').bold = True
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    p = doc.add_paragraph()
    p.add_run('{{receiving_party}}')
    doc.add_paragraph()
    doc.add_paragraph('By: _________________________')
    doc.add_paragraph('Name:')
    doc.add_paragraph('Title:')
    doc.add_paragraph('Date:')
    
    # Save the document
    doc.save(output_path)
    print(f'✓ Created {output_path}')
    print('')
    print('Template includes the following variables:')
    print('  - {{effective_date}}')
    print('  - {{disclosing_party}}')
    print('  - {{receiving_party}}')
    print('  - {{term_months}}')
    print('  - {{survival_months}}')
    print('  - {{governing_law}}')
    print('')
    print(f'File saved to: {os.path.abspath(output_path)}')
    return output_path

if __name__ == '__main__':
    output_file = sys.argv[1] if len(sys.argv) > 1 else 'templates/template_standard_mutual_nda.docx'
    # Ensure templates directory exists
    os.makedirs(os.path.dirname(output_file) if os.path.dirname(output_file) else 'templates', exist_ok=True)
    create_nda_template(output_file)

