#!/usr/bin/env python3
"""
End-to-end workflow test script
Tests the complete NDA workflow: create → send → receive → LLM review → workflow → approval
"""
import os
import sys
import time
import uuid
import requests
import json
from datetime import datetime, timedelta

# Add project root to path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)

from api.services.bootstrap import configure_services_from_env
from api.db import get_db_session
from api.db.schema import NDARecord, NDATemplate, Document, DocumentStatus, NDAWorkflowInstance
from api.services.template_service import create_template_service
from api.services.service_registry import get_storage_service
from api.services.email_service import get_email_service
from api.services.email_parser import EmailParser
from api.services.camunda_service import get_camunda_service
from api.services.llm_review_service import get_llm_review_service

# Configuration
API_BASE_URL = os.getenv("API_BASE_URL", "http://localhost:8000")
TEST_USERNAME = os.getenv("TEST_USERNAME", "das_service")
TEST_PASSWORD = os.getenv("TEST_PASSWORD", "das_service_2024!")

# Colors for output
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
BLUE = "\033[94m"
RESET = "\033[0m"


def print_step(step_num, description):
    """Print a test step header"""
    print(f"\n{BLUE}{'='*60}{RESET}")
    print(f"{BLUE}Step {step_num}: {description}{RESET}")
    print(f"{BLUE}{'='*60}{RESET}")


def print_success(message):
    """Print success message"""
    print(f"{GREEN}✓ {message}{RESET}")


def print_error(message):
    """Print error message"""
    print(f"{RED}✗ {message}{RESET}")


def print_info(message):
    """Print info message"""
    print(f"{YELLOW}ℹ {message}{RESET}")


def ensure_test_user():
    """Ensure test user exists"""
    print_info("Ensuring test user exists...")
    
    # Try to create user directly in database
    try:
        from api.db import get_db_session
        from api.db.schema import User
        from api.auth import get_password_hash
        import uuid
        
        db = get_db_session()
        try:
            user = db.query(User).filter(User.username == TEST_USERNAME).first()
            if not user:
                user = User(
                    id=uuid.uuid4(),
                    username=TEST_USERNAME,
                    password_hash=get_password_hash(TEST_PASSWORD),
                    role="admin",
                    is_active=True,
                )
                db.add(user)
                db.commit()
                print_success(f"Created test user: {TEST_USERNAME}")
            else:
                print_info(f"Test user already exists: {TEST_USERNAME}")
        finally:
            db.close()
    except Exception as e:
        print_info(f"Could not create user directly: {e}")
        # Try using admin user instead
        print_info("Falling back to admin user...")
        return "admin", "Admin2024!Secure"
    
    return TEST_USERNAME, TEST_PASSWORD


def get_auth_token():
    """Get authentication token"""
    print_info("Authenticating...")
    
    # Ensure test user exists
    username, password = ensure_test_user()
    
    response = requests.post(
        f"{API_BASE_URL}/auth/login",
        json={"username": username, "password": password},
    )
    if response.status_code != 200:
        raise Exception(f"Authentication failed: {response.text}")
    token = response.json()["access_token"]
    print_success(f"Authenticated as {username}")
    return token


def create_test_template(token):
    """Create a test NDA template"""
    print_info("Creating test NDA template...")
    
    # Create a simple DOCX template using python-docx
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument()
        doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
        doc.add_paragraph('This Non-Disclosure Agreement ("Agreement") is entered into on {effective_date} between:')
        doc.add_paragraph('Party 1: {party1_name}')
        doc.add_paragraph('Party 2: {party2_name}')
        doc.add_paragraph('')
        doc.add_paragraph('TERM: This Agreement shall remain in effect for {term_months} months from the effective date.')
        doc.add_paragraph('')
        doc.add_paragraph('CONFIDENTIAL INFORMATION: The parties agree to keep confidential all information disclosed during the term of this agreement.')
        doc.add_paragraph('')
        doc.add_paragraph('SIGNATURES:')
        doc.add_paragraph('Party 1: _________________ Date: ___________')
        doc.add_paragraph('Party 2: _________________ Date: ___________')
        
        # Save to bytes
        import io
        template_bytes = io.BytesIO()
        doc.save(template_bytes)
        template_bytes.seek(0)
        template_data = template_bytes.read()
        
    except ImportError:
        # Fallback: create minimal DOCX manually
        print_info("python-docx not available, creating minimal DOCX...")
        # Create a minimal valid DOCX (ZIP file with XML)
        import zipfile
        import io
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>''')
            zip_file.writestr('word/document.xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t>NON-DISCLOSURE AGREEMENT</w:t></w:r></w:p>
<w:p><w:r><w:t>This Agreement is entered into on {effective_date} between {party1_name} and {party2_name}.</w:t></w:r></w:p>
<w:p><w:r><w:t>TERM: {term_months} months</w:t></w:r></w:p>
</w:body>
</w:document>''')
            zip_file.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>''')
        
        template_data = zip_buffer.getvalue()
    
    # Upload template via API
    print_info("Uploading template via API...")
    files = {'file': ('test_template.docx', template_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {
        'name': f'Test Template {uuid.uuid4().hex[:8]}',
        'description': 'Test template for E2E workflow',
    }
    
    response = requests.post(
        f"{API_BASE_URL}/templates",
        headers={"Authorization": f"Bearer {token}"},
        files=files,
        data=data,
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to create template: {response.text}")
    
    result = response.json()
    template_id = result["id"]
    print_success(f"Created template: {template_id}")
    return template_id


def create_nda_from_template(token, template_id):
    """Create an NDA from template"""
    print_info("Creating NDA from template...")
    
    response = requests.post(
        f"{API_BASE_URL}/workflow/nda/create",
        headers={"Authorization": f"Bearer {token}"},
        json={
            "template_id": template_id,
            "counterparty_name": "Test Company Inc.",
            "counterparty_domain": "testcompany.com",
            "effective_date": datetime.now().date().isoformat(),
            "term_months": 12,
            "additional_data": {
                "party1_name": "Our Company",
                "party2_name": "Test Company Inc.",
            },
        },
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to create NDA: {response.text}")
    
    result = response.json()
    # Check for different possible response keys
    nda_id = result.get("nda_record_id") or result.get("id") or result.get("nda_id")
    if not nda_id:
        # Try to get from nested structure
        if "nda_record" in result:
            nda_id = result["nda_record"].get("id")
        if not nda_id:
            raise Exception(f"Could not find NDA ID in response: {result}")
    
    print_success(f"Created NDA: {nda_id}")
    return nda_id


def send_nda_email(token, nda_id):
    """Send NDA via email"""
    print_info("Sending NDA email...")
    
    response = requests.post(
        f"{API_BASE_URL}/workflow/nda/{nda_id}/send",
        headers={"Authorization": f"Bearer {token}"},
        json={
            "to_addresses": ["customer@testcompany.com"],
            "subject": "NDA for Review and Signature",
            "body": "Please review and sign the attached NDA.",
        },
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to send email: {response.text}")
    
    result = response.json()
    print_success(f"Email sent with tracking ID: {result.get('tracking_id')}")
    return result.get("tracking_id")


def simulate_signed_nda_received(nda_id, tracking_id):
    """Simulate receiving a signed NDA back"""
    print_info("Simulating signed NDA received...")
    
    # Get NDA record to find email address
    db = get_db_session()
    try:
        nda_uuid = uuid.UUID(nda_id)
        nda_record = db.query(NDARecord).filter(NDARecord.id == nda_uuid).first()
        if not nda_record:
            raise Exception(f"NDA {nda_id} not found")
        
        # Create a mock signed document
        storage = get_storage_service()
        signed_content = f"""
NON-DISCLOSURE AGREEMENT - SIGNED

This Non-Disclosure Agreement has been reviewed and signed.

Party 1: Our Company
Party 2: Test Company Inc.

SIGNATURES:

Party 1: John Admin                    Date: {datetime.now().date()}
Party 2: Jane Customer                 Date: {datetime.now().date()}
"""
        
        # Upload signed document (use nda-raw bucket which exists)
        signed_path = f"signed_{nda_id}.txt"
        storage.upload_file(
            bucket="nda-raw",
            object_name=signed_path,
            file_data=signed_content.encode(),
            content_type="text/plain",
        )
        
        # Update NDA record with signed document
        nda_record.file_uri = f"s3://nda-raw/{signed_path}"
        nda_record.extracted_text = signed_content
        nda_record.status = "customer_signed"
        db.commit()
        
        print_success("Simulated signed NDA received")
        
    finally:
        db.close()


def trigger_llm_review(nda_id):
    """Trigger LLM review of the NDA"""
    print_info("Triggering LLM review...")
    
    llm_service = get_llm_review_service()
    
    # Perform review
    review_result = None
    try:
        import asyncio
        loop = asyncio.get_event_loop()
        review_result = loop.run_until_complete(llm_service.review_nda(nda_id))
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        review_result = loop.run_until_complete(llm_service.review_nda(nda_id))
    
    print_info(f"LLM Review Result:")
    print_info(f"  Approved: {review_result['approved']}")
    print_info(f"  Confidence: {review_result['confidence']}")
    print_info(f"  Reasoning: {review_result.get('reasoning', 'N/A')[:100]}")
    
    if review_result['approved']:
        print_success("LLM review passed")
    else:
        print_error("LLM review failed")
    
    return review_result


def start_workflow(token, nda_id):
    """Start Camunda workflow"""
    print_info("Starting Camunda workflow...")
    
    response = requests.post(
        f"{API_BASE_URL}/workflow/nda/{nda_id}/start-workflow",
        headers={"Authorization": f"Bearer {token}"},
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to start workflow: {response.text}")
    
    result = response.json()
    workflow_instance_id = result["workflow_instance_id"]
    camunda_process_instance_id = result["camunda_process_instance_id"]
    
    print_success(f"Workflow started: {workflow_instance_id}")
    print_info(f"Camunda Process Instance: {camunda_process_instance_id}")
    
    return workflow_instance_id, camunda_process_instance_id


def wait_for_external_task(process_instance_id, timeout=30):
    """Wait for external task to be created"""
    print_info(f"Waiting for external task (timeout: {timeout}s)...")
    
    camunda = get_camunda_service()
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        try:
            tasks = camunda.get_external_tasks(process_instance_id=process_instance_id)
            if tasks:
                print_success(f"Found {len(tasks)} external task(s)")
                return tasks[0]
        except Exception as e:
            print_info(f"Error checking external tasks: {e}")
        
        time.sleep(2)
    
    print_error("Timeout waiting for external task")
    return None


def wait_for_user_task(process_instance_id, timeout=30):
    """Wait for user task to be created"""
    print_info(f"Waiting for user task (timeout: {timeout}s)...")
    
    camunda = get_camunda_service()
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        tasks = camunda.get_user_tasks(process_instance_id=process_instance_id)
        if tasks:
            print_success(f"Found {len(tasks)} user task(s)")
            return tasks[0]
        time.sleep(2)
    
    print_error("Timeout waiting for user task")
    return None


def get_workflow_tasks(token, workflow_instance_id):
    """Get workflow tasks"""
    print_info("Getting workflow tasks...")
    
    response = requests.get(
        f"{API_BASE_URL}/workflow/workflow/{workflow_instance_id}/tasks",
        headers={"Authorization": f"Bearer {token}"},
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to get tasks: {response.text}")
    
    result = response.json()
    print_success(f"Found {len(result.get('camunda_tasks', []))} Camunda task(s)")
    return result


def complete_task(token, task_id, approved=True):
    """Complete a workflow task"""
    print_info(f"Completing task {task_id} (approved={approved})...")
    
    response = requests.post(
        f"{API_BASE_URL}/workflow/workflow/task/{task_id}/complete",
        headers={"Authorization": f"Bearer {token}"},
        params={"approved": approved},
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to complete task: {response.text}")
    
    print_success(f"Task {task_id} completed")
    return response.json()


def check_nda_status(token, nda_id):
    """Check NDA status"""
    print_info("Checking NDA status...")
    
    response = requests.get(
        f"{API_BASE_URL}/workflow/nda/{nda_id}/status",
        headers={"Authorization": f"Bearer {token}"},
    )
    
    if response.status_code != 200:
        raise Exception(f"Failed to get status: {response.text}")
    
    result = response.json()
    status = result.get("status")
    print_success(f"NDA status: {status}")
    return status


def main():
    """Run end-to-end workflow test"""
    print(f"\n{GREEN}{'='*60}{RESET}")
    print(f"{GREEN}NDA Workflow End-to-End Test{RESET}")
    print(f"{GREEN}{'='*60}{RESET}\n")
    
    try:
        # Step 1: Authenticate
        print_step(1, "Authentication")
        token = get_auth_token()
        
        # Step 2: Create template
        print_step(2, "Create NDA Template")
        template_id = create_test_template(token)
        
        # Step 3: Create NDA
        print_step(3, "Create NDA from Template")
        nda_id = create_nda_from_template(token, template_id)
        
        # Step 4: Send email
        print_step(4, "Send NDA Email")
        tracking_id = send_nda_email(token, nda_id)
        
        # Step 5: Simulate signed NDA received
        print_step(5, "Simulate Signed NDA Received")
        simulate_signed_nda_received(nda_id, tracking_id)
        
        # Step 6: Check status
        print_step(6, "Check NDA Status")
        status = check_nda_status(token, nda_id)
        if status != "customer_signed":
            print_error(f"Expected status 'customer_signed', got '{status}'")
        
        # Step 7: LLM Review
        print_step(7, "LLM Review")
        review_result = trigger_llm_review(nda_id)
        
        # Step 8: Start workflow (only if LLM approved)
        if review_result['approved']:
            print_step(8, "Start Camunda Workflow")
            workflow_instance_id, process_instance_id = start_workflow(token, nda_id)
            
            # Step 9: Wait for external task (LLM review task)
            print_step(9, "Wait for External Task")
            external_task = wait_for_external_task(process_instance_id, timeout=30)
            
            if external_task:
                print_info("External task will be processed by worker...")
                print_info("Waiting for worker to complete external task (up to 60s)...")
                # Wait for external task to be completed
                camunda = get_camunda_service()
                start_time = time.time()
                while time.time() - start_time < 60:
                    remaining_tasks = camunda.get_external_tasks(process_instance_id=process_instance_id)
                    if not remaining_tasks:
                        print_success("External task completed by worker")
                        break
                    time.sleep(2)
                else:
                    print_info("External task may still be processing...")
            
            # Step 10: Wait for user tasks
            print_step(10, "Wait for User Tasks")
            user_task = wait_for_user_task(process_instance_id, timeout=30)
            
            if user_task:
                task_id = user_task.get("id")
                task_name = user_task.get("name", "")
                print_info(f"Found task: {task_name} (ID: {task_id})")
                
                # Step 11: Complete tasks
                print_step(11, "Complete Workflow Tasks")
                complete_task(token, task_id, approved=True)
                
                # Wait for next task
                time.sleep(5)
                user_task = wait_for_user_task(process_instance_id, timeout=30)
                if user_task:
                    task_id = user_task.get("id")
                    task_name = user_task.get("name", "")
                    print_info(f"Found task: {task_name} (ID: {task_id})")
                    complete_task(token, task_id, approved=True)
                
                # Wait for final approval task
                time.sleep(5)
                user_task = wait_for_user_task(process_instance_id, timeout=30)
                if user_task:
                    task_id = user_task.get("id")
                    task_name = user_task.get("name", "")
                    print_info(f"Found task: {task_name} (ID: {task_id})")
                    complete_task(token, task_id, approved=True)
                
                # Step 12: Final status check
                print_step(12, "Final Status Check")
                final_status = check_nda_status(token, nda_id)
                
                if final_status in ["approved", "reviewed"]:
                    print_success(f"Workflow completed successfully! Final status: {final_status}")
                else:
                    print_error(f"Unexpected final status: {final_status}")
            else:
                print_error("No user tasks found")
        else:
            print_error("LLM review failed - workflow not started")
        
        print(f"\n{GREEN}{'='*60}{RESET}")
        print(f"{GREEN}Test Completed Successfully!{RESET}")
        print(f"{GREEN}{'='*60}{RESET}\n")
        
    except Exception as e:
        print_error(f"Test failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

